from __future__ import annotations

from dataclasses import dataclass
from functools import lru_cache
from io import BytesIO
from pathlib import Path


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".png", ".jpg", ".jpeg", ".webp", ".txt"}
IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp"}


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    file_type: str
    content: str
    extraction_method: str


def extract_text_from_file(filename: str, data: bytes) -> ExtractedDocument:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Nieobslugiwany typ pliku. Obslugiwane rozszerzenia: {supported}.")

    if extension == ".pdf":
        content = _extract_pdf_text(data)
        method = "pypdf"
    elif extension == ".docx":
        content = _extract_docx_text(data)
        method = "python-docx"
    elif extension in IMAGE_EXTENSIONS:
        content = _extract_image_text(data)
        method = "easyocr"
    else:
        content = _extract_plain_text(data)
        method = "plain-text"

    content = _clean_text(content)
    if len(content.split()) < 10:
        raise ValueError("Nie udalo sie wyciagnac wystarczajacej ilosci tekstu z pliku.")

    return ExtractedDocument(
        filename=filename,
        file_type=extension.removeprefix("."),
        content=content,
        extraction_method=method,
    )


def _extract_pdf_text(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n\n".join(pages)


def _extract_docx_text(data: bytes) -> str:
    from docx import Document

    document = Document(BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text)
    return "\n\n".join(paragraphs + table_cells)


def _extract_image_text(data: bytes) -> str:
    try:
        import numpy as np
        from PIL import Image, ImageOps

        reader = _get_easyocr_reader()
    except ImportError as exc:
        raise RuntimeError("OCR wymaga bibliotek Pillow, numpy i easyocr.") from exc

    image = ImageOps.exif_transpose(Image.open(BytesIO(data))).convert("RGB")
    fragments = reader.readtext(np.array(image), detail=0, paragraph=True)
    return "\n".join(fragment.strip() for fragment in fragments if fragment.strip())


@lru_cache(maxsize=1)
def _get_easyocr_reader():
    import easyocr

    return easyocr.Reader(["pl", "en"], gpu=False, verbose=False)


def _extract_plain_text(data: bytes) -> str:
    for encoding in ("utf-8", "windows-1250", "iso-8859-2"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def _clean_text(text: str) -> str:
    lines = [line.strip() for line in text.replace("\x00", " ").splitlines()]
    return "\n".join(line for line in lines if line)
