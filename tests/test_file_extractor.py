import sys
from io import BytesIO
from types import SimpleNamespace

import pytest
from docx import Document
from PIL import Image
from pypdf import PdfWriter
from pypdf.generic import DecodedStreamObject, DictionaryObject, NameObject

from app.services.file_extractor import extract_text_from_file


def test_extract_plain_text_file() -> None:
    extracted = extract_text_from_file(
        "article.txt",
        "Ministerstwo opublikowało raport z danymi. Tekst zawiera liczby, źródła i spokojny opis sprawy.".encode(
            "utf-8"
        ),
    )

    assert extracted.file_type == "txt"
    assert extracted.extraction_method == "plain-text"
    assert "raport" in extracted.content


def test_extract_docx_file() -> None:
    document = Document()
    document.add_paragraph("Agencja opublikowała raport o jakości wody.")
    document.add_paragraph("Tekst zawiera dane, daty, ekspertów i linki do źródeł publicznych.")
    buffer = BytesIO()
    document.save(buffer)

    extracted = extract_text_from_file("article.docx", buffer.getvalue())

    assert extracted.file_type == "docx"
    assert extracted.extraction_method == "python-docx"
    assert "jakości wody" in extracted.content


def test_extract_pdf_file() -> None:
    writer = PdfWriter()
    page = writer.add_blank_page(width=300, height=200)
    font = DictionaryObject(
        {
            NameObject("/Type"): NameObject("/Font"),
            NameObject("/Subtype"): NameObject("/Type1"),
            NameObject("/BaseFont"): NameObject("/Helvetica"),
        }
    )
    font_ref = writer._add_object(font)
    page[NameObject("/Resources")] = DictionaryObject(
        {NameObject("/Font"): DictionaryObject({NameObject("/F1"): font_ref})}
    )
    stream = DecodedStreamObject()
    stream.set_data(b"BT /F1 12 Tf 50 150 Td (Agency published a report with data and experts for verification.) Tj ET")
    page[NameObject("/Contents")] = writer._add_object(stream)
    buffer = BytesIO()
    writer.write(buffer)

    extracted = extract_text_from_file("article.pdf", buffer.getvalue())

    assert extracted.file_type == "pdf"
    assert extracted.extraction_method == "pypdf"
    assert "Agency published a report" in extracted.content


def test_extract_image_file_with_easyocr(monkeypatch: pytest.MonkeyPatch) -> None:
    class FakeReader:
        def __init__(self, languages: list[str], gpu: bool, verbose: bool) -> None:
            assert languages == ["pl", "en"]
            assert gpu is False
            assert verbose is False

        def readtext(self, image: object, detail: int, paragraph: bool) -> list[str]:
            assert detail == 0
            assert paragraph is True
            return [
                "Agency published a report with data and experts.",
                "The article links official research sources for verification.",
            ]

    monkeypatch.setitem(sys.modules, "easyocr", SimpleNamespace(Reader=FakeReader))
    monkeypatch.setitem(sys.modules, "numpy", SimpleNamespace(array=lambda image: image))

    image = Image.new("RGB", (320, 120), color="white")
    buffer = BytesIO()
    image.save(buffer, format="PNG")

    extracted = extract_text_from_file("screenshot.png", buffer.getvalue())

    assert extracted.file_type == "png"
    assert extracted.extraction_method == "easyocr"
    assert "official research sources" in extracted.content


def test_rejects_unsupported_file_type() -> None:
    with pytest.raises(ValueError, match="Nieobslugiwany typ pliku"):
        extract_text_from_file("archive.zip", b"not an article")
